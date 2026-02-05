import pdfplumber
import docx
import spacy
import re
import datetime

# ---------- SPAcY LAZY LOADER ----------

def get_nlp():
    global nlp
    if "nlp" not in globals():
        nlp = spacy.load("en_core_web_sm")
    return nlp

# ---------- TEXT EXTRACTION ----------

def extract_pdf_text(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = ''
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_docx_text(file_path):
    doc = docx.Document(file_path)

    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return '\n'.join(full_text)

# ---------- NAME EXTRACTION ----------

def extract_names(text):
    lines = text.split('\n')

    # Strategy 1: First 5 lines heuristic
    for line in lines[:5]:
        line = line.strip()
        if 5 < len(line) < 50:
            if line.replace(' ', '').replace('.', '').isalpha():
                if line.isupper() or line.istitle():
                    return [line]

    # Strategy 2: spaCy NER
    top_section = text[:500]
    nlp_model = get_nlp()
    doc = nlp_model(top_section)

    person_entities = [ent.text for ent in doc.ents if ent.label_ == 'PERSON']

    exclude_terms = [
        'Tamil Nadu', 'India', 'Chennai', 'Mumbai', 'Delhi',
        'Bangalore', 'Wells Fargo', 'Infosys', 'College'
    ]

    filtered_names = [n for n in person_entities if n not in exclude_terms]
    if filtered_names:
        return [filtered_names[0]]

    # Strategy 3: Name near email
    email_match = re.search(r'([A-Za-z\s]+)\s*[\|\-]?\s*[\w.]+@', text[:500])
    if email_match:
        potential_name = email_match.group(1).strip()
        if 5 < len(potential_name) < 50:
            return [potential_name]

    return []

# ---------- BASIC INFO ----------

def extract_email(text):
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    match = re.search(email_regex, text)
    return match.group(0) if match else None


def extract_phone(text):
    patterns = [
        r'\+91[-\s]?\d{5}[-\s]?\d{5}',
        r'\b[6-9]\d{9}\b',
        r'\+91[-\s]?\d{10}'
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)

    return None

# ---------- SKILLS ----------

def extract_skills(text):
    skills = [
        "Python", "Machine Learning", "SQL", "Java", "C++",
        "TensorFlow", "PyTorch", "Deep Learning", "NLP",
        "Data Analysis", "scikit-learn", "Pandas", "NumPy"
    ]
    return [s for s in skills if s.lower() in text.lower()]

# ---------- EDUCATION ----------

def extract_education(text):
    section_headers = ["education", "academic", "qualification", "educational background"]
    lines = text.split('\n')

    start_idx = -1
    for i, line in enumerate(lines):
        if any(h in line.lower() for h in section_headers):
            start_idx = i
            break

    if start_idx != -1:
        education_section = '\n'.join(lines[start_idx:start_idx + 15])
    else:
        education_section = text

    keywords = [
        "B.Tech", "M.Tech", "B.E.", "MCA", "BSc", "B.Sc",
        "MSc", "M.Sc", "MBA", "BCA", "Diploma",
        "Bachelor", "Master", "PhD",
        "Computer Science", "Information Technology", "Mechanical Engineering"
    ]

    found = [k for k in keywords if k.lower() in education_section.lower()]
    return list(set(found))

# ---------- EXPERIENCE ----------

def extract_experience_years(text):
    direct_pattern = r'(\d+)\+?\s*years?\s*(?:of)?\s*(?:experience|exp)?'
    matches = re.findall(direct_pattern, text.lower())
    if matches:
        return max(int(m) for m in matches)

    date_ranges = re.findall(r'(20\d{2})\s*[-–]\s*(20\d{2}|present|current)', text.lower())
    if date_ranges:
        current_year = datetime.datetime.now().year
        total_years = 0
        for start, end in date_ranges:
            start_year = int(start)
            end_year = current_year if end in ['present', 'current'] else int(end)
            total_years += (end_year - start_year)
        return total_years if total_years > 0 else None

    return None

# ---------- JOB TITLES ----------

def extract_job_titles(text):
    job_titles = [
        "Software Engineer", "Senior Software Engineer", "Lead Software Engineer",
        "Software Developer", "Full Stack Developer", "Backend Developer",
        "Frontend Developer", "DevOps Engineer", "Cloud Engineer",
        "Data Scientist", "Data Analyst", "Data Engineer",
        "Machine Learning Engineer", "AI Engineer",
        "Business Analyst", "Project Manager", "Product Manager",
        "Technical Lead", "QA Engineer"
    ]

    found = []
    text_lower = text.lower()

    for title in job_titles:
        if title.lower() in text_lower:
            found.append(title)

    seen = set()
    result = []
    for t in found:
        if t not in seen:
            seen.add(t)
            result.append(t)

    return result[:3]

# ---------- MAIN PARSER ----------

def parse_resume(file_path):
    if file_path.endswith(".pdf"):
        text = extract_pdf_text(file_path)
    elif file_path.endswith(".docx"):
        text = extract_docx_text(file_path)
    else:
        return {
            "name": [],
            "email": "Unsupported format",
            "phone": None,
            "skills": [],
            "education": [],
            "experience_years": None,
            "job_titles": []
        }

    return {
        "name": extract_names(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "education": extract_education(text),
        "experience_years": extract_experience_years(text),
        "job_titles": extract_job_titles(text)
    }
